import docx
import sys

docx_path = r'C:\Users\majidmohamed\Microsoft\MAI - Strategy - 11. MSI strategy\1. Ad hoc\0. Maj Working Folder\3. Adhoc\State of the Union - AI Training Infra\State of the Union - AI Training Infra.docx'

try:
    doc = docx.Document(docx_path)

    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            print(' | '.join(cells))

except Exception as e:
    print(f"Error: {e}", file=sys.stderr)
